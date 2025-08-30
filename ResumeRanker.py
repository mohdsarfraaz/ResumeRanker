# 📄 Resume Ranker - Streamlit App
# --------------------------------------------------

import streamlit as st
import pandas as pd
import PyPDF2
import docx
import os
import re
import base64
import matplotlib.pyplot as plt

from datetime import datetime
from rapidfuzz import fuzz
from sentence_transformers import SentenceTransformer, util


st.set_page_config(layout="wide")

# =========================
# 🎨 Custom Header Styling
# =========================
st.markdown(
    """
    <style>
    .header-container {
        text-align: center;
        margin-bottom: 30px;
    }

    /* Banner Image with 3D effect */
    .header-container img {
        width: 85%;
        max-height: 260px;
        object-fit: cover;
        border-radius: 18px;
        margin-bottom: 20px;
        box-shadow: 0 12px 25px rgba(0, 0, 0, 0.3), 
                    0 8px 15px rgba(0, 0, 0, 0.2) inset;
        transition: transform 0.4s ease, box-shadow 0.4s ease;
    }
    .header-container img:hover {
        transform: scale(1.03) rotateX(5deg);
        box-shadow: 0 18px 30px rgba(0, 0, 0, 0.35);
    }

    /* Title with 3D text shadow */
    .main-title {
        font-size: 46px;
        font-weight: 900;
        color: #2E86C1;
        margin: 10px 0;
        font-family: 'Trebuchet MS', sans-serif;
        text-shadow: 3px 3px 0px #a5c9f7, 6px 6px 0px #d0e4ff;
    }

    /* Subtitle italic with depth */
    .subtitle {
        font-size: 22px;
        color: #333;
        margin-bottom: 25px;
        font-style: italic;
        font-weight: 500;
        text-shadow: 1px 1px 3px rgba(0,0,0,0.15);
    }

    /* Info box with 3D glowing style */
    .info-box {
        text-align: center;
        font-size: 18px;
        padding: 14px 22px;
        border-radius: 14px;
        width: 70%;
        margin: auto;
        font-family: 'Segoe UI', sans-serif;
        background: linear-gradient(145deg, #e8f5e9, #c8efce);
        color: #1b5e20;
        font-weight: 600;
        box-shadow: 4px 4px 10px rgba(0,0,0,0.2),
                    -4px -4px 10px rgba(255,255,255,0.8);
        transition: transform 0.3s ease;
    }
    .info-box:hover {
        transform: scale(1.05);
        box-shadow: 6px 6px 14px rgba(0,0,0,0.25),
                    -6px -6px 14px rgba(255,255,255,0.9);
    }
    </style>

    <div class="header-container">
        <h1 class="main-title">📄 Resume Ranker</h1>
        <div class="subtitle">Smart Resume Shortlisting System for HR Professionals</div>
        <img src="https://img.freepik.com/premium-photo/finding-right-candidate_968957-19679.jpg" alt="HR Banner">
        <div class="info-box">
            💡 Great teams are built by hiring the right talent. Let’s make it faster & smarter!
        </div>
    </div>
    """,
    unsafe_allow_html=True
)


# =========================
# 📌 Model Loading
# =========================
@st.cache_resource
def load_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

model = load_model()


# =========================
# 📌 Helper Functions
# =========================
@st.cache_data
def extract_text_from_path(file_path):
    """Extract text from a PDF or DOCX file path."""
    try:
        text = ""
        if file_path.endswith(".pdf"):
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    if page.extract_text():
                        text += page.extract_text() + "\n"

        elif file_path.endswith(".docx"):
            with open(file_path, "rb") as f:
                document = docx.Document(f)
                for para in document.paragraphs:
                    text += para.text + "\n"

        return text.strip()
    except Exception as e:
        st.error(f"Error processing {os.path.basename(file_path)}: {str(e)}")
        return ""


def extract_experience(text):
    """Extract total years of experience from job date ranges."""
    clean_text = text.replace("’", "'").replace("–", "-").replace("—", "-")
    date_patterns = re.findall(
        r"([A-Za-z]{3,9})[\' ]?(\d{2,4})\s*[-to]+\s*([A-Za-z]{3,9})[\' ]?(\d{2,4}|date|present)",
        clean_text,
        re.IGNORECASE,
    )

    total_months = 0
    for start_month, start_year, end_month, end_year in date_patterns:
        try:
            start_year = int("20" + start_year if len(start_year) == 2 else start_year)
            end_year = (
                datetime.now().year
                if end_year.lower() in ["date", "present"]
                else int("20" + end_year if len(end_year) == 2 else end_year)
            )

            start_date = datetime.strptime(f"{start_month} {start_year}", "%b %Y")
            try:
                end_date = datetime.strptime(f"{end_month} {end_year}", "%b %Y")
            except ValueError:
                end_date = datetime.strptime(f"{end_month} {end_year}", "%B %Y")

            months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
            if months > 0:
                total_months += months
        except Exception:
            continue

    return round(total_months / 12, 1)


def calculate_similarity(jd_text, resume_text):
    """Semantic similarity between JD and resume text."""
    if not jd_text.strip() or not resume_text.strip():
        return 0
    jd_embedding = model.encode(jd_text, convert_to_tensor=True)
    resume_embedding = model.encode(resume_text, convert_to_tensor=True)
    similarity = util.cos_sim(jd_embedding, resume_embedding).item()
    return round(similarity * 100, 2)


def match_skills(required_skills, resume_text, threshold=80):
    """Match required skills with fuzzy matching."""
    matched = []
    resume_text_lower = resume_text.lower()

    for skill in required_skills:
        skill_lower = skill.lower()
        if re.search(r"\b" + re.escape(skill_lower) + r"\b", resume_text_lower):
            matched.append(skill)
        elif fuzz.partial_ratio(skill_lower, resume_text_lower) >= threshold:
            matched.append(skill)

    return list(set(matched))


# =========================
# 📑 Job Description Input
# =========================
st.set_page_config(page_title="📄 Resume Ranker", layout="wide")

# ===============================
# 🚀 Custom CSS
# ===============================
st.markdown(
    """
    <style>
    /* Center main content but keep right panel visible */
    .main-container {
        display: flex;
        flex-direction: row;
        justify-content: space-between;
    }

    /* Sticky right panel */
    .right-panel {
        position: fixed;
        top: 70px;
        right: 10px;
        width: 280px;
        background: #f9f9f9;
        border-radius: 12px;
        padding: 16px;
        box-shadow: 4px 4px 12px rgba(0,0,0,0.1);
        z-index: 999;
        overflow-y: auto;
        max-height: 85vh;
    }

    /* Sponsor cards */
    .sponsor-box {
        background: white;
        border-radius: 10px;
        padding: 10px;
        margin-bottom: 12px;
        text-align: center;
        box-shadow: 2px 2px 8px rgba(0,0,0,0.1);
        font-weight: 500;
    }

    /* Make share buttons inline */
    .share-buttons a {
        margin: 0 8px;
        text-decoration: none;
        font-size: 20px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ===============================
# 📝 Job Description Section
# ===============================
st.subheader("📑 Job Description")


col1, col2 = st.columns([2, 1])  # left JD, right controls

with col2:
    jd_mode = st.radio("Choose input method:", ["✍️ Paste JD", "📂 Upload JD File"])

with col1:
    if jd_mode == "✍️ Paste JD":
        jd_text = st.text_area("Paste Job Description Here", height=200)

    elif jd_mode == "📂 Upload JD File":
        uploaded_jd = st.file_uploader("Upload JD (TXT, DOCX, PDF)", type=["txt", "docx", "pdf"])
        if uploaded_jd is not None:
            if uploaded_jd.name.endswith(".txt"):
                jd_text = uploaded_jd.read().decode("utf-8")
            elif uploaded_jd.name.endswith(".docx"):
                doc = docx.Document(uploaded_jd)
                jd_text = "\n".join([para.text for para in doc.paragraphs])
            elif uploaded_jd.name.endswith(".pdf"):
                pdf_reader = PyPDF2.PdfReader(uploaded_jd)
                jd_text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
            # Preview
            if jd_text.strip():
                st.success("✅ Job Description loaded successfully!")
                st.text_area("Preview JD", jd_text[:1500], height=200)

# =========================
# 🔎 Filters
# =========================
col1, col2 = st.columns(2)
with col1:
    min_score = st.slider("Minimum Match Score (%)", 0, 100, 50)
    min_exp = st.slider("Minimum Experience (years)", 0, 20, 2)
with col2:
    skills_filter = st.text_input("Required skills (comma-separated)", "Python, SQL, Machine Learning")
    location_filter = st.text_input("Preferred Location (optional)")

mode = st.radio("Choose Input Mode:", ["📂 Folder Mode (Local)", "☁️ Upload Mode (Public)"])


# =========================
# 📂 Resume Processing
# =========================
results = []
resume_files_paths = []

if mode == "📂 Folder Mode (Local)":
    resume_folder = "resumes_to_process"
    if not os.path.exists(resume_folder):
        os.makedirs(resume_folder)
        st.warning(f"📂 Created a folder named '{resume_folder}'. Please add your resumes there.")

    if st.button("⚡ Process Resumes (Folder)", type="primary", use_container_width=True):
        resume_files_paths = [
            os.path.join(resume_folder, f)
            for f in os.listdir(resume_folder)
            if f.endswith((".pdf", ".docx"))
        ]

elif mode == "☁️ Upload Mode (Public)":
    uploaded_files = st.file_uploader(
        "📤 Upload Resumes (PDF/DOCX)", 
        type=["pdf", "docx"], 
        accept_multiple_files=True
    )
    if uploaded_files and st.button("⚡ Process Resumes (Uploads)", type="primary", use_container_width=True):
        upload_folder = "uploaded_resumes"
        os.makedirs(upload_folder, exist_ok=True)

        for uploaded_file in uploaded_files:
            resume_path = os.path.join(upload_folder, uploaded_file.name)
            with open(resume_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            resume_files_paths.append(resume_path)

        st.success(f"✅ {len(resume_files_paths)} resumes saved in '{upload_folder}' folder")


# =========================
# 📊 Processing & Results
# =========================
if resume_files_paths:
    if not jd_text.strip():
        st.warning("⚠️ Please provide a Job Description.")
    else:
        with st.spinner("Processing resumes..."):
            required_skills = [s.strip().lower() for s in skills_filter.split(",") if s.strip()]

            for file_path in resume_files_paths:
                file_name = os.path.basename(file_path)

                text = extract_text_from_path(file_path)
                if not text:
                    continue

                score = calculate_similarity(jd_text, text)
                experience = extract_experience(text)
                matched_skills = match_skills(required_skills, text)
                location_match = not location_filter.strip() or (location_filter.strip().lower() in text.lower())

                is_qualified = (
                    score >= min_score
                    and experience >= min_exp
                    and location_match
                    and (len(matched_skills) > 0 if required_skills else True)
                )

                results.append({
                    "Resume": file_name,
                    "Score (%)": round(score, 2),
                    "Experience (yrs)": round(experience, 1),
                    "Qualified": "✅ Yes" if is_qualified else "❌ No",
                    "Matched Skills": ", ".join(list(set(matched_skills))) or "None",
                    "Location Match": "✅ Yes" if location_match else "❌ No",
                    "File Path": file_path
                })

        if results:
            df = pd.DataFrame(results).sort_values(by="Score (%)", ascending=False).reset_index(drop=True)
            df.index += 1

            df["Score (%)"] = df["Score (%)"].map(lambda x: f"{x:.2f}")
            df["Experience (yrs)"] = df["Experience (yrs)"].map(lambda x: f"{x:.1f}")

            # Add download links
            def make_download_link(file_path):
                try:
                    with open(file_path, "rb") as f:
                        data = f.read()
                    b64 = base64.b64encode(data).decode()
                    file_name = os.path.basename(file_path)
                    return f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}">📂 Download</a>'
                except Exception as e:
                    return f"❌ Error: {str(e)}"

            df["Download Resume"] = [make_download_link(path) for path in df["File Path"]]
            if "File Path" in df.columns:
                df = df.drop(columns=["File Path"])

            st.subheader("📊 Candidate Screening Results")
            st.dataframe(df)

            # ✅ Qualified / ❌ Unqualified
            def render_table(dataframe, title):
                st.subheader(title)
                if dataframe.empty:
                    st.info("No data available.")
                else:
                    styled_df = dataframe.style.set_properties(**{"text-align": "left"})
                    st.markdown(styled_df.to_html(escape=False), unsafe_allow_html=True)

            render_table(df[df["Qualified"] == "✅ Yes"], "✅ Qualified Candidates")
            render_table(df[df["Qualified"] == "❌ No"], "❌ Unqualified Candidates")

            # 📥 Export CSV
            csv = df.to_csv(index_label="Rank").encode("utf-8")
            st.download_button("📥 Download Full Results (CSV)", csv, "resume_ranking.csv", "text/csv")

            # ----------------- DASHBOARD -----------------
            st.markdown("## 📊 HR Resume Screening Dashboard")

            # Metrics
            col1, col2, col3 = st.columns(3)
            with col1: st.metric("Total Resumes", len(df))
            with col2: st.metric("Qualified", (df["Qualified"] == "✅ Yes").sum())
            with col3: st.metric("Unqualified", (df["Qualified"] == "❌ No").sum())

            # Score Distribution
            
            with col1: 
                st.markdown("### 🎯 Match Score Distribution")
                fig, ax = plt.subplots()
                df["Score (%)"].astype(float).plot(kind="hist", bins=10, ax=ax)
                ax.set_xlabel("Match Score (%)")
                st.pyplot(fig)

            # Experience Distribution
            with col2:
                st.markdown("### ⏳ Experience (Years) Distribution")
                fig, ax = plt.subplots()
                df["Experience (yrs)"].astype(float).plot(kind="hist", bins=10, ax=ax)
                ax.set_xlabel("Experience (Years)")
                st.pyplot(fig)

            # Word Cloud for Skills
            try:
                from wordcloud import WordCloud
                all_skills = " ".join(df["Matched Skills"].dropna())
                if all_skills.strip():
                    wordcloud = WordCloud(width=800, height=400, background_color="white").generate(all_skills)
                    st.markdown("### 🛠️ Common Matched Skills")
                    fig, ax = plt.subplots(figsize=(8, 4))
                    ax.imshow(wordcloud, interpolation="bilinear")
                    ax.axis("off")
                    st.pyplot(fig)
            except ImportError:
                st.info("Install `wordcloud` to see skill cloud visualization.")


# =========================
# 📌 Footer
# =========================
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; padding: 15px; font-size: 18px; font-family: cursive; color: grey;">
        ✍️ Developed by <b>Md Sarfraaz</b>
    </div>
    """,
    unsafe_allow_html=True
)
# --- Right Interactive Panel (new) ---
with st.sidebar:
    st.header("✨ Interact")
    st.text_area("💬 Leave a comment")
    st.button("📤 Share")
    st.text_input("📝 Feedback")
    st.header("🤝 Sponsors")
    st.markdown("Sponsor logos or ads here...")
